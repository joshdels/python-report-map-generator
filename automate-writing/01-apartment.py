from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

# -----------------------------
# CONFIG
# -----------------------------
LESSOR = {
    "name": "MONICA S. DELA CRUZ",
    "address": "Brgy. New Pandan, Davao City, Philippines",
    "id": "SSD ID No.0X-123123123-X"
}

LESSEE = {
    "name": "JOSHUA S. DE LEON",
    "address": "Brgy. Visayan Village, Tagum City, Philippines",
    "id": "PRC REG. NO.X00232322X"
}

PROPERTY = {
    "name": "Lucky 7, Doors Apartment",
    "location": "Purok 5, Camelia Homes Subd,",
    "lot_num": "500-G-32-A-C",
    "license": "TCT No. 100-X01505000"
}

RENT = {
    "cost_num": 6000,
    "cost_string": "six thousand",
    "start_date": "October 4, 2025"
}

HOUSE_RULES = [
    "That the LESSOR shall not be liable for any damage by caused by or any damage arising from the failure of the water supply and or electricity installment or the bursting leaking or running of any wash pipes, rain water or water closet and the the LESSOR shall not be liable",
    "No Pets Allowed inside the apartment",
    "The LESSEE shall not sublease the apartment without the written consent of the LESSOR;",
    "No illegal activities shall be conducted within the premises"
]

INDENT = Pt(36)

# -----------------------------
# HELPERS
# -----------------------------
def set_heading(doc, text, level=1, align=WD_ALIGN_PARAGRAPH.CENTER):
    try:
        h = doc.add_heading(text, level)
        h.alignment = align
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
        return h
    except Exception as e:
        print(f"⚠️ Failed to set heading '{text}': {e}")
        return None

def add_bold(paragraph, text):
    try:
        run = paragraph.add_run(text)
        run.bold = True
        return run
    except Exception as e:
        print(f"⚠️ Failed to add bold text '{text}': {e}")
        return None

def signature_block(cell, name, role, id_no):
    try:
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(name).bold = True

        cell.add_paragraph(role).alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.add_paragraph(id_no).alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        print(f"⚠️ Failed to create signature block for '{name}': {e}")

# -----------------------------
# MAIN DOCUMENT CREATION
# -----------------------------
try:
    document = Document()

    # HEADING
    set_heading(document, "CONTRACT OF LEASE")
    set_heading(document, "KNOW ALL MEN BY THESE PRESENTS", level=2, align=WD_ALIGN_PARAGRAPH.LEFT)

    # MINI TEXT
    p = document.add_paragraph("This ")
    add_bold(p, "CONTRACT OF LEASES")
    p.add_run(" made and entered into by and between:")
    p.first_line_indent = INDENT
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph()
    add_bold(p, LESSOR.get("name", ""))
    p.add_run(", of legal age, married and a resident of ")
    p.add_run(LESSOR.get("address", ""))
    p.add_run(" now and hereinafter called the")
    add_bold(p, "LESSOR; ")
    p.add_run("  -and--")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = document.add_paragraph()
    add_bold(p, LESSEE.get("name", ""))
    p.add_run(f", also a Filipino, of legal age, single and a resident of {LESSEE.get('address','')}")
    p.add_run(" now and hereinafter known as the ")
    add_bold(p, "LESSEE; ")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # WITNESSETH
    withnesseth = set_heading(document, "WITHNESSETH:")
    if withnesseth:
        withnesseth.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # MAIN BODY
    p = document.add_paragraph("That herein the LESSOR ")
    add_bold(p, LESSOR.get("name", ""))
    p.add_run(" is known as the ")
    add_bold(p, "OWNER ")
    p.add_run(f" of {PROPERTY.get('name', '')} at {PROPERTY.get('location', '')}")
    p.add_run(f" covered by {PROPERTY.get('license', '')}.")
    p.paragraph_format.first_line_indent = INDENT
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # HOUSE RULES
    set_heading(document, "HOUSE RULES:", level=3, align=WD_ALIGN_PARAGRAPH.LEFT)
    for i in HOUSE_RULES:
        p = document.add_paragraph(i, "List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # AGREEMENT
    p = document.add_paragraph("That either parties may terminate the lease by giving at least seven (7) days written notice. The ")
    add_bold(p, "LESSEE")
    p.add_run(" must vacate the premises in good and clean condition.")    
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = INDENT

    document.add_page_break()

    # TABLE SIGNATURE
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    signature_block(table.cell(0, 0), LESSOR.get("name",""), "Lessor", LESSOR.get("id",""))
    signature_block(table.cell(0, 1), LESSEE.get("name",""), "Lessee", LESSEE.get("id",""))

    # NOTARY SECTION
    p = document.add_paragraph("")
    p = document.add_paragraph("IN WITNESS WHEREOF, the parties hereto set their hands this ______________, at ______________, Philippines.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = INDENT

    p = document.add_paragraph(
        "BEFORE ME, a Notary Public for and the city of Davao/Province of Davao, personally appeared the above name with their valid ID's written below their names, ")
    p.add_run("which known to me to the very same persons who executed the foregoing instrument and they acknowledge to me the same is their free and voluntary deed.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = INDENT

    # SAVE FILE
    output_path = "../results/letter.docx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    document.save(output_path)
    print(f"✅ File saved at {output_path}!")

except Exception as e:
    print(f"❌ Failed to create document: {e}")
